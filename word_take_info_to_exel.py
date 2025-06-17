import docx
import pandas as pd
from openpyxl import load_workbook
import os
import glob
import re
import numpy as np

def get_text(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_tables(docx_path):
    """
    Extract all tables from a Word file as lists of lists (rows of cells).
    Returns a list of tables, each table is a list of rows, each row is a list of cell texts.
    """
    doc = docx.Document(docx_path)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables

def extract_meeting_info(text, tables=None):
    def extract_after(label, text, multiline=False):
        pattern = rf"{re.escape(label)}\s*(.*)"
        m = re.search(pattern, text)
        return m.group(1).strip() if m else ''
    info = {}
    # If tables are provided, extract from the first table
    if tables and len(tables) > 0:
        table = tables[0]
        def get_cells_below(keyword):
            values = []
            found = False
            for row in table:
                if found:
                    if row[0].strip().endswith(":") and row[0].strip() != keyword:
                        break
                    if row[0].strip() and row[0].strip() != keyword and not row[0].strip().endswith(":"):
                        values.append(row[0].strip())
                    if len(row) > 1 and row[1].strip():
                        values.append(row[1].strip())
                if row[0].strip() == keyword:
                    found = True
            return values
        info['Голова комітету'] = ", ".join(get_cells_below('Голова комітету:'))
        info['Заступник Голови комітету'] = ", ".join(get_cells_below('Заступник Голови комітету:'))
        info['Секретар комітету'] = ", ".join(get_cells_below('Секретар комітету (без права голосу):'))
        info['Члени комітету'] = get_cells_below('Члени комітету:')
        
    info['Відсутні'] = extract_after('Відсутні:', text)
    info['Запрошені'] = extract_after('Запрошені:', text)
    info['Кворум'] = extract_after('Кворум:', text)
    info['Порядок прийняття рішень'] = extract_after('Порядок прийняття рішень:', text)
    return info

def extract_questions(text):
    questions = []
    # Extract agenda topics
    agenda_match = re.search(r'ПОРЯДОК ДЕННИЙ:(.*?)(?:Питання 1:|$)', text, re.DOTALL)
    agenda = []
    if agenda_match:
        agenda = [line.strip(" .\t") for line in agenda_match.group(1).split('\n') if line.strip()]
    # Find all "Питання N:" blocks
    question_blocks = list(re.finditer(r'Питання\s*\d+:(.*?)(?=Питання\s*\d+:|ПІБ|\Z)', text, re.DOTALL))
    for idx, match in enumerate(question_blocks):
        block = match.group(1)
        # Speaker
        speaker_match = re.search(r'Виступив:(.*?)Голосували:', block, re.DOTALL)
        speaker = speaker_match.group(1).strip() if speaker_match else ''
        # Votes
        votes_match = re.search(r'Голосували:(.*?)Вирішили:', block, re.DOTALL)
        votes = {'за': '', 'проти': '', 'утримались': ''}
        if votes_match:
            for v in votes:
                m = re.search(rf'{v}\s*-\s*(\S+)', votes_match.group(1))
                if m:
                    votes[v] = m.group(1)
        # Decision
        decision_match = re.search(r'Вирішили:(.*)', block)
        decision = decision_match.group(1).strip() if decision_match else ''
        
        # Use agenda topic if available, else fallback to first line
        if idx < len(agenda):
            question_text = agenda[idx]
        else:
            question_text = block.split('\n')[0].strip()
        questions.append({
            'question text': question_text,
            'speaker': speaker,
            'votes': votes,
            'decision': decision
        })
    return questions

def extract_vote_table(text):
    # Find the table at the end (ПІБ, за, проти, утримались, примітка)
    table = []
    table_match = re.search(r'ПІБ\s+за\s+проти\s+утримались\s+примітка(.*)', text, re.DOTALL)
    if not table_match:
        return table
    lines = table_match.group(1).split('\n')
    current_name = None
    for line in lines:
        if not line.strip():
            continue
        # If line contains only a name
        if re.match(r'^[А-ЯІЇЄҐ][а-яіїєґ]+\s+[А-ЯІЇЄҐ][а-яіїєґ]+\s+[А-ЯІЇЄҐ][а-яіїєґ]+$', line.strip()):
            current_name = line.strip()
            continue
        # If line contains vote marks
        cells = line.strip().split('\t')
        if len(cells) >= 4 and current_name:
            table.append({
                'ПІБ': current_name,
                'за': cells[0].strip(),
                'проти': cells[1].strip(),
                'утримались': cells[2].strip(),
                'примітка': cells[3].strip() if len(cells) > 3 else ''
            })
            current_name = None
        elif len(cells) >= 4:
            # Sometimes name and votes are on the same line
            table.append({
                'ПІБ': cells[0].strip(),
                'за': cells[1].strip(),
                'проти': cells[2].strip(),
                'утримались': cells[3].strip(),
                'примітка': cells[4].strip() if len(cells) > 4 else ''
            })
    return table

def extract_doc_number_and_date(text):
    # Extract document number (handles underscores and spaces)
    number_match = re.search(r'ПРОТОКОЛ\s*№\s*_*(\S+)_*', text, re.IGNORECASE)
    number = number_match.group(1).strip('_') if number_match else ''
    # Extract date (handles underscores, spaces, and Ukrainian months)
    date_match = re.search(
        r'«\s*_?(\d{1,2})_?\s*»\s*_+([а-яА-Яіїєґ]+)_+\s*(\d{4})\s*року', text)
    month_map = {
        'січня': '01', 'лютого': '02', 'березня': '03', 'квітня': '04',
        'травня': '05', 'червня': '06', 'липня': '07', 'серпня': '08',
        'вересня': '09', 'жовтня': '10', 'листопада': '11', 'грудня': '12'
    }
    if date_match:
        day = date_match.group(1)
        month_ua = date_match.group(2).strip('_').lower()
        year = date_match.group(3)
        month = month_map.get(month_ua, '')
        if month:
            date = f"{int(day):02d}.{month}.{year}"
        else:
            date = f"{day} {month_ua} {year}"
    else:
        date = ''
    return number, date

if __name__ == "__main__":
    folder = os.path.dirname(os.path.abspath(__file__))
    file_list = glob.glob(os.path.join(folder, '*.docx'))
    all_rows = []

    for file_path in file_list:
        filename = os.path.basename(file_path)
        text = get_text(file_path)
        tables = extract_tables(file_path)
        info = extract_meeting_info(text, tables)
        questions = extract_questions(text)
        vote_table = extract_vote_table(text)
        number, date = extract_doc_number_and_date(text)
        
        # For each agenda item (Питання), create a block of rows
        for q_idx, q in enumerate(questions):
            # First row: meeting info + first member
            first_member = vote_table[0] if vote_table else {'ПІБ': '', 'за': '', 'проти': '', 'утримались': '', 'примітка': ''}
            meeting_row = {
                'Файл': filename,
                'Дата': date,
                'Номер протоколу': number,
                'Голова комітету': info['Голова комітету'],
                'Заступник Голови комітету': info['Заступник Голови комітету'],
                'Члени комітету': ", ".join(info['Члени комітету']) if isinstance(info['Члени комітету'], list) else info['Члени комітету'],
                'Секретар комітету': info['Секретар комітету'],
                'Відсутні': info['Відсутні'],
                'Запрошені': info['Запрошені'],
                'Кворум': info['Кворум'],
                'Порядок прийняття рішень': info['Порядок прийняття рішень'],
                'ПОРЯДОК ДЕННИЙ': q['question text'],
                'Виступив:': q['speaker'],
                'Голосували: за': q['votes']['за'],
                'Голосували: проти': q['votes']['проти'],
                'Голосували: утримались': q['votes']['утримались'],
                'Вирішили:': q['decision'],
                'ПІБ': first_member['ПІБ'],
                'за': first_member['за'],
                'проти': first_member['проти'],
                'утримались': first_member['утримались'],
                'примітка': first_member['примітка']
            }
            all_rows.append(meeting_row)
            for member in vote_table[1:]:
                member_row = {k: np.nan for k in meeting_row}
                member_row['Члени комітету'] = member['ПІБ']
                member_row['ПІБ'] = member['ПІБ']
                member_row['за'] = member['за']
                member_row['проти'] = member['проти']
                member_row['утримались'] = member['утримались']
                member_row['примітка'] = member['примітка']
                all_rows.append(member_row)

        # Process the last table and add its rows to all_rows, mapping columns by header if possible
        if tables:
            last_table = tables[-1]
            if last_table:
                headers = last_table[0]
                for row in last_table[1:]:
                    row_dict = {}
                    for i, header in enumerate(headers):
                        if i < len(row):
                            row_dict[header] = row[i]
                        else:
                            row_dict[header] = ""
                    all_rows.append(row_dict)
    
    # Create a DataFrame and save to Excel
    df = pd.DataFrame(all_rows)
    df.to_excel('TOTAL_PROTOKOL.xlsx', index=False)
    
    wb = load_workbook('TOTAL_PROTOKOL.xlsx')
    ws = wb.active
    
    for column_cells in ws.columns:
        length = max(len(str(cell.value)) if cell.value else 0 for cell in column_cells)
        col_letter = column_cells[0].column_letter
        ws.column_dimensions[col_letter].width = length + 2

    wb.save('TOTAL_PROTOKOL.xlsx')
    
    print("Extraction complete. Results saved to TOTAL_PROTOKOL.xlsx.")
    input("Press Enter to exit...")